import streamlit as st
from docx import Document
from docx.shared import RGBColor
import zipfile
import re
import tempfile

def processar_docx(uploaded_files, texto_referencia, cor_rgb):
    arquivos_processados = []

    for uploaded_file in uploaded_files:
        doc = Document(uploaded_file)
        alterado = False

        for p in doc.paragraphs:
            if texto_referencia in p.text and '{' in p.text and '}' in p.text:
                match = re.search(r'\{(.+?)\}', p.text)
                if match:
                    dentro_chave = match.group(1)
                    antes = p.text.split('{')[0]
                    depois = p.text.split('}')[1] if '}' in p.text else ''

                    p.clear()
                    p.add_run(antes)

                    run_colorido = p.add_run(dentro_chave)
                    run_colorido.font.color.rgb = cor_rgb

                    p.add_run(depois)
                    alterado = True

        if alterado:
            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(tmp_file.name)
            arquivos_processados.append((uploaded_file.name, tmp_file.name))

    return arquivos_processados

# 🖼️ Interface Streamlit
st.title("📝 Editor DOCX: Destacar texto entre chaves")

st.markdown("""
Faça upload de um ou mais arquivos `.docx`.  
Informe o texto-chave que deve aparecer no parágrafo onde será destacado o texto entre `{}`.  
Escolha a cor em hexadecimal (ex: `#FF0000` para vermelho).
""")

uploaded_files = st.file_uploader("📁 Selecione arquivos DOCX", type="docx", accept_multiple_files=True)
texto_referencia = st.text_input("🔍 Texto-chave", value="será na data:")
cor_hex = st.color_picker("🎨 Escolha a cor de destaque", value="#FF0000")

# 🔄 Conversão de HEX para RGBColor
cor_hex_clean = cor_hex.lstrip("#")
try:
    r, g, b = tuple(int(cor_hex_clean[i:i+2], 16) for i in (0, 2, 4))
    cor_rgb = RGBColor(r, g, b)
except:
    st.error("❌ Código de cor inválido. Use formato como: #FF0000")
    st.stop()

if st.button("✅ Processar arquivos"):
    if not uploaded_files:
        st.warning("Por favor, envie pelo menos um arquivo `.docx`.")
    else:
        resultado = processar_docx(uploaded_files, texto_referencia, cor_rgb)
        if resultado:
            zip_path = tempfile.NamedTemporaryFile(delete=False, suffix=".zip").name
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for nome_arquivo, caminho_temp in resultado:
                    zipf.write(caminho_temp, arcname=nome_arquivo)
            st.success(f"{len(resultado)} arquivo(s) processado(s) com sucesso!")
            with open(zip_path, "rb") as f:
                st.download_button("📥 Baixar arquivos editados (.zip)", f.read(), file_name="documentos_editados.zip")
        else:
            st.info("Nenhum arquivo continha o texto-chave com texto entre chaves `{}` para editar.")
